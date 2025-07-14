import streamlit as st
import requests
import io
from docx import Document
from fpdf import FPDF
import unicodedata
import json

def remove_accents(text):
    return ''.join(
        c for c in unicodedata.normalize('NFD', text)
        if unicodedata.category(c) != 'Mn'
    )

# Display logo at the top
st.image("../logo.png", width=120)

st.title("Générateur de Contenu de Formation IA")

# User input
with st.form("course_form"):
    topic = st.text_input("Sujet du cours", "")
    level = st.selectbox("Niveau", ["débutant", "intermédiaire", "avancé"])
    model_options = [
        ("Llama 4 Maverick", "meta-llama/llama-4-maverick"),
        ("GPT-4o-mini", "openai/gpt-4o-mini"),
        ("Gemini 2.0 Flash", "google/gemini-2.0-flash-001"),
        ("QwQ 32B", "qwen/qwq-32b:free"),
        ("Deepseek R1 0528 Qwen3 8B", "deepseek/deepseek-r1-0528-qwen3-8b:free")
    ]
    model_display_names = [name for name, _ in model_options]
    model_selected = st.selectbox("Modèle IA", model_display_names)
    model_id = dict(model_options)[model_selected]
    col1, spacer, col2 = st.columns([1, 0.05, 1])
    with col1:
        submitted = st.form_submit_button("Générer le cours", use_container_width=True)
    with col2:
        submitted_qcm = st.form_submit_button("Générer QCM", use_container_width=True)
    

if submitted:
    with st.spinner("Génération en cours..."):
        try:
            response = requests.post(
                "http://127.0.0.1:8000/generate-course",
                json={"topic": topic, "level": level, "model": model_id},
                timeout=30
            )
            response.raise_for_status()
            data = response.json()
            st.session_state["course_data"] = data
            st.success("Cours généré !")
        except requests.Timeout:
            st.error("La requête a expiré. Veuillez réessayer plus tard.")
        except requests.RequestException as e:
            st.error(f"Erreur lors de la génération : {e}")
        except Exception as e:
            st.error(f"Erreur inattendue : {e}")
if submitted_qcm:
    with st.spinner("Génération du QCM en cours..."):
        try:
            response = requests.post(
                "http://127.0.0.1:8000/generate-qcm",
                json={"topic": topic, "level": level, "model": model_id},
                timeout=30
            )
            response.raise_for_status()
            qcm_data = response.json()["qcm"]
            st.session_state["qcm_data"] = qcm_data  # Already a Python list/dict
            st.success("QCM généré !")
        except requests.Timeout:
            st.error("La requête a expiré. Veuillez réessayer plus tard.")
        except requests.RequestException as e:
            st.error(f"Erreur lors de la génération du QCM : {e}")
        except Exception as e:
            st.error(f"Erreur inattendue : {e}")


if "course_data" in st.session_state:
    data = st.session_state["course_data"]
    st.subheader(data["title"])
    st.markdown(f"**Plan :**\n{data['outline']}")
    st.markdown(f"**Résumé :**\n{data['summary']}")
    st.markdown(f"**Contenu brut :**\n{data['raw_content']}")

    export_col1, export_col2 = st.columns(2)
    with export_col1:
        if st.button("Exporter en Word (.docx)"):
            doc = Document()
            doc.add_heading(data["title"], 0)
            doc.add_heading("Plan", level=1)
            doc.add_paragraph(data["outline"])
            doc.add_heading("Résumé", level=1)
            doc.add_paragraph(data["summary"])
            doc.add_heading("Contenu brut", level=1)
            doc.add_paragraph(data["raw_content"])
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            st.download_button(
                label="Télécharger le Word",
                data=buf,
                file_name="cours_formaia.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        if st.button("Exporter en PowerPoint (.pptx)"):
            with st.spinner("Génération du PowerPoint en cours..."):
                try:
                    response = requests.post(
                        "http://127.0.0.1:8000/generate-course-pptx",
                        json={
                            "topic": topic,
                            "level": level,
                            "model": model_id
                        },
                        timeout=60
                    )
                    response.raise_for_status()
                    pptx_bytes = response.content
                    st.download_button(
                        label="Télécharger le PowerPoint",
                        data=pptx_bytes,
                        file_name="cours_formaia.pptx",
                        mime="application/vnd.openxmlformats-officedocument.presentationml.presentation"
                    )
                except requests.Timeout:
                    st.error("La requête a expiré. Veuillez réessayer plus tard.")
                except requests.RequestException as e:
                    st.error(f"Erreur lors de la génération du PowerPoint : {e}")
                except Exception as e:
                    st.error(f"Erreur inattendue : {e}")
    with export_col2:
        if st.button("Exporter en PDF"):
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=16)
            pdf.cell(0, 10, remove_accents(data["title"]), ln=True, align="C")
            pdf.set_font("Arial", size=12)
            pdf.cell(0, 10, "Plan:", ln=True)
            pdf.multi_cell(0, 10, remove_accents(data["outline"]))
            pdf.cell(0, 10, "Résumé:", ln=True)
            pdf.multi_cell(0, 10, remove_accents(data["summary"]))
            pdf.cell(0, 10, "Contenu brut:", ln=True)
            pdf.multi_cell(0, 10, remove_accents(data["raw_content"]))
            pdf_output = pdf.output(dest='S').encode('latin1')
            st.download_button(
                label="Télécharger le PDF",
                data=pdf_output,
                file_name="cours_formaia.pdf",
                mime="application/pdf"
            )
if "qcm_data" in st.session_state:
    data = st.session_state["qcm_data"]
    st.subheader("QCM")
    st.markdown("### Questions :")

    for i, question in enumerate(data):
        st.write(f"**Question {i + 1}:** {question['question']}")
        st.radio(
            "Choix :",
            options=question['choices'],
            key=f"qcm_{i}",
            index=None
        )

    qcm_button = st.button("Soumettre le QCM")

    if qcm_button:
        score = 0
        st.markdown("---")
        st.markdown("### Résultats :")

        for i, question in enumerate(data):
            user_answer = st.session_state.get(f"qcm_{i}", None)
            correct_letter = question["answer"]
            correct_text = next(
                (choice for choice in question["choices"] if choice.startswith(correct_letter)), None
            )

            if user_answer == correct_text:
                st.success(f"✅ Question {i + 1}: Correct")
                score += 1
            else:
                st.error(f"❌ Question {i + 1}: Incorrect (Bonne réponse : {correct_text})")

        st.markdown("---")
        st.markdown(f"### Score final : **{score} / {len(data)}**")


    